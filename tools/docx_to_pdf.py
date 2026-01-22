"""Simple script to extract text from all .docx files under docs/ and
write corresponding .pdf files next to them.

Usage:
    py tools\docx_to_pdf.py

This uses python-docx (pip package name: python-docx).
"""
import sys
from pathlib import Path

try:
    from docx import Document
except Exception as e:
    print("Missing python-docx. Install with: py -m pip install python-docx")
    raise


def docx_to_text(path: Path) -> str:
    doc = Document(path)
    parts = []
    for para in doc.paragraphs:
        parts.append(para.text)
    # Also try tables
    for table in doc.tables:
        for row in table.rows:
            row_text = '\t'.join(cell.text for cell in row.cells)
            parts.append(row_text)
    return "\n".join(parts)


def main():
    root = Path(__file__).resolve().parents[1] / "docs"
    if not root.exists():
        print(f"Docs folder not found at {root}")
        sys.exit(1)

    docx_files = list(root.rglob('*.docx'))
    if not docx_files:
        print("No .docx files found under docs/")
        return

    for f in docx_files:
        try:
            pdf = docx_to_text(f)
            out = f.with_suffix('.pdf')
            out.write_text(pdf, encoding='utf-8')
            print(f"Wrote: {out}")
        except Exception as e:
            print(f"Failed to process {f}: {e}")


if __name__ == '__main__':
    main()
